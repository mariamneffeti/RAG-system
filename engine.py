"""
RAG Engine — Document loader, chunking, embeddings, FAISS vector DB, Groq LLM.
"""

import os
import re
import json
import pickle
import hashlib
from pathlib import Path
from dataclasses import dataclass, field
from typing import Iterator

import numpy as np
import faiss
from sentence_transformers import SentenceTransformer


#  Data Structures 

@dataclass
class Document:
    content: str
    metadata: dict = field(default_factory=dict)
    doc_id: str = ""

    def __post_init__(self):
        if not self.doc_id:
            self.doc_id = hashlib.md5(self.content.encode()).hexdigest()[:12]


@dataclass
class Chunk:
    text: str
    chunk_id: str
    doc_id: str
    chunk_index: int
    metadata: dict = field(default_factory=dict)


@dataclass
class SearchResult:
    chunk: Chunk
    score: float


#  Document Loaders 

def load_txt(path: str) -> Document:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    return Document(content=content, metadata={"source": path, "type": "txt"})


def load_pdf(path: str) -> Document:
    try:
        import PyPDF2
        text = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text.append(page.extract_text() or "")
    except ImportError:
        raise ImportError("pip install PyPDF2")
    return Document(content="\n".join(text), metadata={"source": path, "type": "pdf"})


def load_docx(path: str) -> Document:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(path)
        content = "\n".join(p.text for p in doc.paragraphs)
    except ImportError:
        raise ImportError("pip install python-docx")
    return Document(content=content, metadata={"source": path, "type": "docx"})


def load_file(path: str) -> Document:
    ext = Path(path).suffix.lower()
    loaders = {".txt": load_txt, ".md": load_txt, ".pdf": load_pdf, ".docx": load_docx}
    loader = loaders.get(ext)
    if not loader:
        raise ValueError(f"Unsupported file type: {ext}")
    return loader(path)


def load_text(text: str, source_name: str = "inline") -> Document:
    return Document(content=text, metadata={"source": source_name, "type": "text"})


def load_directory(dir_path: str) -> list[Document]:
    docs = []
    for p in Path(dir_path).rglob("*"):
        if p.suffix.lower() in (".txt", ".md", ".pdf", ".docx") and p.is_file():
            try:
                docs.append(load_file(str(p)))
            except Exception as e:
                print(f"[WARN] Skipping {p}: {e}")
    return docs


#  Chunking 

def chunk_document(doc: Document, chunk_size: int = 512, chunk_overlap: int = 64) -> list[Chunk]:
    text = doc.content.strip()
    if not text:
        return []
    sentences = re.split(r"(?<=[.!?])\s+", text)
    segments  = []
    current, current_len = [], 0
    for sent in sentences:
        if current_len + len(sent) > chunk_size and current:
            segments.append(" ".join(current))
            current, current_len = [], 0
        current.append(sent)
        current_len += len(sent)
    if current:
        segments.append(" ".join(current))

    return [
        Chunk(
            text=seg,
            chunk_id=f"{doc.doc_id}-{i}",
            doc_id=doc.doc_id,
            chunk_index=i,
            metadata={**doc.metadata, "chunk_index": i, "total_chunks": len(segments)},
        )
        for i, seg in enumerate(segments)
    ]


#  Embeddings 

class EmbeddingModel:
    def __init__(self, model_name: str = "paraphrase-MiniLM-L3-v2"):
        print(f"[EMB] Loading: {model_name}")
        self.model = SentenceTransformer(model_name)
        self.dim   = self.model.get_sentence_embedding_dimension()
        print(f"[EMB] Dim: {self.dim}")

    def embed(self, texts: list[str], batch_size: int = 64) -> np.ndarray:
        return self.model.encode(
            texts, batch_size=batch_size, normalize_embeddings=True
        ).astype(np.float32)

    def embed_one(self, text: str) -> np.ndarray:
        return self.embed([text])[0]


#  FAISS Vector Store 

class VectorStore:
    def __init__(self, dim: int):
        self.dim    = dim
        self.index  = faiss.IndexFlatIP(dim)
        self.chunks: list[Chunk] = []

    def add(self, chunks: list[Chunk], embeddings: np.ndarray):
        self.index.add(embeddings)
        self.chunks.extend(chunks)
        print(f"[VS] Total chunks: {len(self.chunks)}")

    def search(self, query_vec: np.ndarray, top_k: int = 5) -> list[SearchResult]:
        if self.index.ntotal == 0:
            return []
        scores, indices = self.index.search(
            query_vec.reshape(1, -1), min(top_k, self.index.ntotal)
        )
        return [
            SearchResult(chunk=self.chunks[idx], score=float(score))
            for score, idx in zip(scores[0], indices[0]) if idx >= 0
        ]

    def save(self, path: str):
        Path(path).mkdir(parents=True, exist_ok=True)
        faiss.write_index(self.index, f"{path}/index.faiss")
        with open(f"{path}/chunks.pkl", "wb") as f:
            pickle.dump(self.chunks, f)
        with open(f"{path}/meta.json", "w") as f:
            json.dump({"dim": self.dim, "total": len(self.chunks)}, f)

    @classmethod
    def load(cls, path: str) -> "VectorStore":
        with open(f"{path}/meta.json") as f:
            meta = json.load(f)
        vs = cls(dim=meta["dim"])
        vs.index = faiss.read_index(f"{path}/index.faiss")
        with open(f"{path}/chunks.pkl", "rb") as f:
            vs.chunks = pickle.load(f)
        print(f"[VS] Loaded {len(vs.chunks)} chunks")
        return vs


#  RAG Pipeline 

GROQ_MODEL  = "llama-3.3-70b-versatile"
SYSTEM_PROMPT = (
    "You are a helpful assistant. Answer questions using ONLY the provided context. "
    "If the answer is not in the context, say so clearly. Be concise and accurate."
)

class RAGPipeline:
    def __init__(
        self,
        model_name: str    = "paraphrase-MiniLM-L3-v2",
        chunk_size: int    = 512,
        chunk_overlap: int = 64,
    ):
        self.embedder      = EmbeddingModel(model_name)
        self.vector_store  = VectorStore(self.embedder.dim)
        self.chunk_size    = chunk_size
        self.chunk_overlap = chunk_overlap
        self.docs: list[Document] = []

    #  Ingestion 

    def ingest(self, docs: list[Document]):
        all_chunks = []
        for doc in docs:
            self.docs.append(doc)
            all_chunks.extend(chunk_document(doc, self.chunk_size, self.chunk_overlap))
        if not all_chunks:
            return
        print(f"[RAG] Embedding {len(all_chunks)} chunks...")
        self.vector_store.add(all_chunks, self.embedder.embed([c.text for c in all_chunks]))

    def ingest_text(self, text: str, source: str = "inline"):
        self.ingest([load_text(text, source)])

    def ingest_file(self, path: str):
        self.ingest([load_file(path)])

    def ingest_directory(self, path: str):
        self.ingest(load_directory(path))

    #  Retrieval 

    def retrieve(self, query: str, top_k: int = 5) -> list[SearchResult]:
        return self.vector_store.search(self.embedder.embed_one(query), top_k=top_k)

    def build_context(self, results: list[SearchResult], max_chars: int = 3000) -> str:
        parts, total = [], 0
        for r in results:
            block = f"[Source: {r.chunk.metadata.get('source','?')} | Score: {r.score:.3f}]\n{r.chunk.text.strip()}"
            if total + len(block) > max_chars:
                break
            parts.append(block)
            total += len(block)
        return "\n\n---\n\n".join(parts)

    #  Streaming Query (Groq) 

    def query_stream(self, question: str, top_k: int = 5, api_key: str = "") -> Iterator[str]:
        results = self.retrieve(question, top_k=top_k)
        context = self.build_context(results)

        if not context:
            yield f"data: {json.dumps('No relevant documents found. Please ingest some documents first.')}\n\n"
            yield "data: [DONE]\n\n"
            return

        # Send sources to the frontend
        yield f"data: __SOURCES__{json.dumps([{'source': r.chunk.metadata.get('source','?'), 'score': round(r.score, 3), 'preview': r.chunk.text[:120]} for r in results])}\n\n"

        # Stream answer from Groq
        try:
            from groq import Groq
            key    = api_key or os.environ.get("GROQ_API_KEY", "")
            client = Groq(api_key=key)
            stream = client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user",   "content": f"Context:\n{context}\n\nQuestion: {question}"},
                ],
                stream=True,
                max_tokens=1024,
            )
            for chunk in stream:
                token = chunk.choices[0].delta.content
                if token:
                    yield f"data: {json.dumps(token)}\n\n"

        except Exception as e:
            yield f"data: {json.dumps(f'Groq error: {str(e)}')}\n\n"

        yield "data: [DONE]\n\n"

    # Persistence 

    def save(self, path: str = "./rag_store"):
        self.vector_store.save(path)

    def load(self, path: str = "./rag_store"):
        self.vector_store = VectorStore.load(path)

    @property
    def stats(self) -> dict:
        return {
            "documents": len(self.docs),
            "chunks":    len(self.vector_store.chunks),
            "model":     GROQ_MODEL,
        }